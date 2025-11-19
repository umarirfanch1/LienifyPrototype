# app.py - Lienify (Streamlit) - Optimized, lazy template load and single-question flow
import streamlit as st
from docx import Document
from datetime import datetime, timedelta
import zipfile
import tempfile
import os
import io

st.set_page_config(page_title="Lienify - Arizona Lien Waiver Generator", layout="centered")

# -----------------------
# Configuration
# -----------------------
ZIP_NAME = "02_Templates-20251119T041237Z-1-001.zip"  # update if needed
STATES = ["Arizona","California","Nevada","Texas","Florida","Georgia","Washington","Oregon","Colorado","Utah","New Mexico","Idaho"]

# Map logical selection -> filename base (note: actual files are .docx inside the zip,
# but your filenames in templates were named with .pdf; we match on base text)
TEMPLATE_MAP = {
    ("Progress", "Yes"): "CONDITIONAL WAIVER AND RELEASE ON PROGRESS PAYMENT.pdf",
    ("Progress", "No"): "UNCONDITIONAL WAIVER AND RELEASE ON PROGRESS PAYMENT.pdf",
    ("Final", "Yes"): "CONDITIONAL WAIVER AND RELEASE ON FINAL PAYMENT.pdf",
    ("Final", "No"): "UNCONDITIONAL WAIVER AND RELEASE ON FINAL PAYMENT.pdf",
}

# -----------------------
# Utility helpers (no heavy IO during UI)
# -----------------------
def human_date(d):
    return d.strftime("%B %d, %Y") if d else ""

def parse_currency_input(s):
    if not s:
        return None
    s = s.strip().replace("$", "").replace(",", "")
    try:
        v = float(s)
        return "${:,.2f}".format(v)
    except:
        return None

def find_template_in_zip(zip_path, expected_name_base, state_folder="arizona"):
    """
    Search inside the ZIP for a file whose filename contains expected_name_base (case-insensitive)
    under the state_folder. Return the zip internal name (path) if found, else None.
    """
    expected_base = expected_name_base.replace(".pdf", "").lower()
    with zipfile.ZipFile(zip_path, "r") as z:
        for fn in z.namelist():
            low = fn.lower()
            # ensure it is within the state folder
            if state_folder.lower() in low and expected_base in os.path.basename(low):
                return fn
    return None

def read_template_bytes_from_zip(zip_path, internal_name):
    with zipfile.ZipFile(zip_path, "r") as z:
        return z.read(internal_name)

def try_open_docx_from_bytes(bts):
    """
    Save bytes to temp .docx and open with python-docx Document
    """
    tf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    try:
        tf.write(bts)
        tf.flush()
        tf.close()
        doc = Document(tf.name)
    finally:
        try:
            os.unlink(tf.name)  # remove temporary file after Document loaded (Document reads into memory)
        except Exception:
            pass
    return doc

def replace_placeholders_in_doc(doc, mapping):
    # Replace paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    # Replace inside tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for k, v in mapping.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)

# -----------------------
# Session state init
# -----------------------
if "step" not in st.session_state:
    st.session_state.step = 0
# flags to prevent repeated auto-advances
for name in ["role_done","payment_type_done","payment_received_done","first_delivery_done","work_through_done"]:
    if name not in st.session_state:
        st.session_state[name] = False

# store field values (start empty if not set)
for key in ["state","role","payment_type","payment_received","first_delivery","work_through_date",
            "OwnerName","ProjectAddress","CustomerName","LienorName","LicenseNumber",
            "PaymentAmount","ExecutionDate","JobNumber","PropertyDescription"]:
    if key not in st.session_state:
        st.session_state[key] = None

# -----------------------
# Small UI helpers
# -----------------------
def render_app_header():
    st.markdown("# Lienify — Lien and Waiver form generator")
    st.markdown("Please fill out required form *")

def nav_buttons(can_back=True, back_target=None, next_label="Next →"):
    cols = st.columns([1,1,1])
    with cols[0]:
        if can_back:
            if st.button("Back"):
                if back_target is None:
                    st.session_state.step = max(0, st.session_state.step - 1)
                else:
                    st.session_state.step = back_target
    with cols[2]:
        # Next button will be handled per-page so we leave here for layout
        pass

# -----------------------
# Step 0 - Welcome
# -----------------------
if st.session_state.step == 0:
    st.markdown("# Welcome to Lienify Waiver and Lien Form Generator")
    if st.button("Select your state"):
        st.session_state.show_state_select = True

    if st.session_state.get("show_state_select", False):
        st.selectbox("Choose state", [""] + STATES, key="state_selector")
        # when state_selector gets set, we observe it below after rendering control

    # if the user chose a state in session
    chosen = st.session_state.get("state_selector", "")
    if chosen:
        st.session_state.state = chosen
        if chosen != "Arizona":
            st.warning("Only Arizona is open for testing right now. Please select Arizona to continue.")
            # user can change selection; we don't advance
        else:
            # Auto advance when Arizona selected
            st.session_state.step = 1

# -----------------------
# Step 1 - Arizona compliance (only shown once, after AZ selected)
# -----------------------
elif st.session_state.step == 1:
    render_app_header()
    st.markdown("**Thanks for selecting Arizona — have a look at compliance check before we proceed**")
    st.info("""
**Arizona compliance summary (quick):**
- Preliminary Notice: must be sent **within 20 days of first delivery**.  
- Conditional waivers bind upon evidence of payment; Unconditional bind upon signing.  
- Waiving lien rights before performing work is illegal.  
- Unlicensed contractors may lose lien rights.  
- Stop notices allowed on private projects (except owner-occupied dwellings).  
- Payment bonds, tenant-as-agent, highway projects and UPL rules may affect lien eligibility.
""")
    cols = st.columns([1,1,1])
    with cols[0]:
        if st.button("Back"):
            st.session_state.step = 0
    with cols[2]:
        if st.button("Yes, I understood. Please proceed"):
            st.session_state.step = 2

# -----------------------
# Steps 2.. : Single-question-per-page flow
# We'll implement auto-advance on selection AND a Next button (disabled until value set).
# -----------------------
else:
    render_app_header()

    # Helper to show a Next button that increments only when clicked
    def show_next(enabled=True, label="Next →"):
        cols = st.columns([1,1,1])
        with cols[2]:
            if enabled:
                if st.button(label):
                    st.session_state.step += 1
            else:
                st.button(label, disabled=True)

    # Page: Role
    if st.session_state.step == 2:
        st.markdown("### 1 of 13 — Your role on this project *")
        role_val = st.selectbox("", ["", "Contractor","Subcontractor","Supplier","Material Provider"], key="role_widget")
        # auto-advance once and only once
        if role_val and not st.session_state.role_done:
            st.session_state.role = role_val
            st.session_state.role_done = True
            st.session_state.step += 1
        # show Next (disabled until selected)
        show_next(enabled=bool(role_val))
        if st.button("Back"):
            st.session_state.step = 1

    # Page: Payment Type
    elif st.session_state.step == 3:
        st.markdown("### 2 of 13 — Is this waiver for a Progress or Final payment? *")
        pt = st.radio("", ("", "Progress", "Final"), key="payment_type_widget")
        if pt and not st.session_state.payment_type_done:
            st.session_state.payment_type = pt
            st.session_state.payment_type_done = True
            st.session_state.step += 1
        show_next(enabled=bool(pt))
        if st.button("Back"):
            st.session_state.step = 2

    # Page: Payment Received
    elif st.session_state.step == 4:
        st.markdown("### 3 of 13 — Has payment been received for this waiver? *")
        pr = st.radio("", ("", "Yes", "No"), key="payment_received_widget")
        if pr and not st.session_state.payment_received_done:
            st.session_state.payment_received = pr
            st.session_state.payment_received_done = True
            st.session_state.step += 1
        show_next(enabled=bool(pr))
        if st.button("Back"):
            st.session_state.step = 3

    # Page: First Delivery Date
    elif st.session_state.step == 5:
        st.markdown("### 4 of 13 — First Delivery Date (required) *")
        fd = st.date_input("", key="first_delivery_widget")
        if fd and not st.session_state.first_delivery_done:
            st.session_state.first_delivery = fd
            st.session_state.first_delivery_done = True
            st.session_state.step += 1
        # show PN calculation if set
        if st.session_state.first_delivery:
            pn_deadline = st.session_state.first_delivery + timedelta(days=20)
            st.info(f"Preliminary Notice deadline = {pn_deadline.strftime('%Y-%m-%d')}")
        show_next(enabled=bool(fd))
        if st.button("Back"):
            st.session_state.step = 4

    # Page: Work Through Date (only required for Progress)
    elif st.session_state.step == 6:
        if st.session_state.payment_type != "Progress":
            # skip automatically
            st.session_state.step += 1
            st.experimental_rerun()
        st.markdown("### 5 of 13 — Work Through Date (required for Progress payments) *")
        wtd = st.date_input("", key="work_through_widget")
        if wtd and not st.session_state.work_through_done:
            st.session_state.work_through_date = wtd
            st.session_state.work_through_done = True
            st.session_state.step += 1
        show_next(enabled=bool(wtd))
        if st.button("Back"):
            st.session_state.step = 5

    # Page: Owner Name
    elif st.session_state.step == 7:
        st.markdown("### 6 of 13 — Owner Name *")
        oname = st.text_input("", key="OwnerName_widget")
        if oname and oname.strip():
            st.session_state.OwnerName = oname.strip()
        show_next(enabled=bool(st.session_state.OwnerName))
        if st.button("Next →"):
            if not st.session_state.OwnerName:
                st.error("Please enter Owner Name.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 6

    # Project Address
    elif st.session_state.step == 8:
        st.markdown("### 7 of 13 — Project / Job Address *")
        addr = st.text_input("", key="ProjectAddress_widget")
        if addr and addr.strip():
            st.session_state.ProjectAddress = addr.strip()
        show_next(enabled=bool(st.session_state.ProjectAddress))
        if st.button("Next →"):
            if not st.session_state.ProjectAddress:
                st.error("Please enter Project Address.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 7

    # Customer / Paying Entity
    elif st.session_state.step == 9:
        st.markdown("### 8 of 13 — Customer / Paying Entity Name *")
        cust = st.text_input("", key="CustomerName_widget")
        if cust and cust.strip():
            st.session_state.CustomerName = cust.strip()
        show_next(enabled=bool(st.session_state.CustomerName))
        if st.button("Next →"):
            if not st.session_state.CustomerName:
                st.error("Please enter Customer Name.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 8

    # Lienor / Contractor
    elif st.session_state.step == 10:
        st.markdown("### 9 of 13 — Lienor / Contractor / Provider Name *")
        lio = st.text_input("", key="LienorName_widget")
        if lio and lio.strip():
            st.session_state.LienorName = lio.strip()
        show_next(enabled=bool(st.session_state.LienorName))
        if st.button("Next →"):
            if not st.session_state.LienorName:
                st.error("Please enter Lienor Name.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 9

    # License Number
    elif st.session_state.step == 11:
        st.markdown("### 10 of 13 — Contractor / Lienor License Number *")
        lic = st.text_input("", key="LicenseNumber_widget")
        if lic and lic.strip():
            st.session_state.LicenseNumber = lic.strip()
        show_next(enabled=bool(st.session_state.LicenseNumber))
        if st.button("Next →"):
            if not st.session_state.LicenseNumber:
                st.error("Please enter License Number.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 10

    # Payment Amount (single field, placeholder shows $0.00)
    elif st.session_state.step == 12:
        st.markdown("### 11 of 13 — Payment Amount *")
        amt = st.text_input("", placeholder="$0.00", key="PaymentAmount_widget")
        parsed = parse_currency_input(amt)
        if parsed:
            st.session_state.PaymentAmount = parsed
        elif amt and not parsed:
            st.error("Payment value not recognized. Use format like $1234.56 or 1234.56")
        show_next(enabled=bool(parsed))
        if st.button("Next →"):
            if not parsed:
                st.error("Please enter a valid payment amount.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 11

    # Execution Date
    elif st.session_state.step == 13:
        st.markdown("### 12 of 13 — Execution Date (when waiver is signed) *")
        ed = st.date_input("", value=datetime.today(), key="ExecutionDate_widget")
        if ed:
            st.session_state.ExecutionDate = ed
        show_next(enabled=bool(st.session_state.ExecutionDate))
        if st.button("Next →"):
            if not st.session_state.ExecutionDate:
                st.error("Please select Execution Date.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 12

    # Job Number
    elif st.session_state.step == 14:
        st.markdown("### 13 of 13 — Job / Project Number *")
        jn = st.text_input("", key="JobNumber_widget")
        if jn and jn.strip():
            st.session_state.JobNumber = jn.strip()
        show_next(enabled=bool(st.session_state.JobNumber))
        if st.button("Next →"):
            if not st.session_state.JobNumber:
                st.error("Please enter Job Number.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 13

    # Property Description (text area) - after JobNumber we go to review
    elif st.session_state.step == 15:
        st.markdown("### Property Description / Legal Description *")
        pd = st.text_area("", height=150, key="PropertyDescription_widget")
        if pd and pd.strip():
            st.session_state.PropertyDescription = pd.strip()
        show_next(enabled=bool(st.session_state.PropertyDescription))
        if st.button("Next →"):
            if not st.session_state.PropertyDescription:
                st.error("Please enter Property Description.")
            else:
                st.session_state.step += 1
        if st.button("Back"):
            st.session_state.step = 14

    # Review & Generate
    elif st.session_state.step == 16:
        st.subheader("Review & Generate")
        st.markdown("Below are the details you provided. If everything looks correct, press **Generate & Download** to create your filled waiver form.")
        # present in natural friendly layout
        st.markdown("---")
        left_col, right_col = st.columns([2,1])
        with left_col:
            st.markdown(f"**Owner:** {st.session_state.OwnerName or ''}")
            st.markdown(f"**Project Address:** {st.session_state.ProjectAddress or ''}")
            st.markdown(f"**Customer / Paying Entity:** {st.session_state.CustomerName or ''}")
            st.markdown(f"**Lienor / Contractor:** {st.session_state.LienorName or ''}")
            st.markdown(f"**License Number:** {st.session_state.LicenseNumber or ''}")
            st.markdown(f"**Payment Amount:** {st.session_state.PaymentAmount or ''}")
            st.markdown(f"**Work Through Date:** {human_date(st.session_state.work_through_date) or 'N/A'}")
            st.markdown(f"**Execution Date:** {human_date(st.session_state.ExecutionDate) or ''}")
            st.markdown(f"**Job / Project Number:** {st.session_state.JobNumber or ''}")
            st.markdown(f"**Property Description:** {st.session_state.PropertyDescription or ''}")
            st.markdown(f"**First Delivery Date:** {human_date(st.session_state.first_delivery) or ''}")
        with right_col:
            # indicate chosen template type (explain)
            ptype = st.session_state.payment_type or "Progress"
            paid = st.session_state.payment_received or "No"
            cond = "Yes" if paid == "No" else "No"
            template_name = TEMPLATE_MAP.get((ptype, cond), "Selected Template")
            st.markdown("**Form Type Selected**")
            st.markdown(f"**{template_name}**")
            st.caption("This form will be populated with the details above and provided as a .docx file for download.")
        st.markdown("---")

        # Buttons
        cols = st.columns([1,1,1])
        with cols[0]:
            if st.button("Back"):
                st.session_state.step = 15
        with cols[2]:
            if st.button("Generate & Download"):
                # Validate required before generating
                required_checks = [
                    ("OwnerName", st.session_state.OwnerName),
                    ("ProjectAddress", st.session_state.ProjectAddress),
                    ("CustomerName", st.session_state.CustomerName),
                    ("LienorName", st.session_state.LienorName),
                    ("LicenseNumber", st.session_state.LicenseNumber),
                    ("PaymentAmount", st.session_state.PaymentAmount),
                    ("ExecutionDate", st.session_state.ExecutionDate),
                    ("JobNumber", st.session_state.JobNumber),
                    ("PropertyDescription", st.session_state.PropertyDescription),
                    ("FirstDelivery", st.session_state.first_delivery),
                ]
                missing = [k for k,v in required_checks if not v]
                # WorkThroughDate required for Progress
                if st.session_state.payment_type == "Progress" and not st.session_state.work_through_date:
                    missing.append("WorkThroughDate (required for Progress)")
                if missing:
                    st.error("Please complete all required fields before generating. Missing: " + ", ".join(missing))
                else:
                    # Begin generation: lazy-load only selected template from ZIP
                    ptype = st.session_state.payment_type
                    paid = st.session_state.payment_received
                    cond = "Yes" if paid == "No" else "No"
                    selected_filename = TEMPLATE_MAP.get((ptype, cond))
                    if not os.path.exists(ZIP_NAME):
                        st.error(f"Template ZIP not found: {ZIP_NAME}")
                    else:
                        internal_name = find_template_in_zip(ZIP_NAME, selected_filename, state_folder="arizona")
                        if not internal_name:
                            st.error("Could not find the selected template inside the ZIP (Arizona folder). Make sure the Arizona folder contains the template.")
                        else:
                            # show spinner / wait message and generate
                            with st.spinner("Please wait… your form is being generated. This may take a few seconds."):
                                try:
                                    template_bytes = read_template_bytes_from_zip(ZIP_NAME, internal_name)
                                    doc = try_open_docx_from_bytes(template_bytes)
                                except Exception as e:
                                    st.error("Failed to open the template as a Word document. Ensure the file inside ZIP is a .docx Word file (even if filename ends with .pdf).")
                                    st.stop()

                                # build mapping for placeholders
                                mapping = {
                                    "{{OwnerName}}": st.session_state.OwnerName or "",
                                    "{{ProjectAddress}}": st.session_state.ProjectAddress or "",
                                    "{{CustomerName}}": st.session_state.CustomerName or "",
                                    "{{LienorName}}": st.session_state.LienorName or "",
                                    "{{LicenseNumber}}": st.session_state.LicenseNumber or "",
                                    "{{PaymentAmount}}": st.session_state.PaymentAmount or "",
                                    "{{WorkThroughDate}}": human_date(st.session_state.work_through_date) if st.session_state.work_through_date else "",
                                    "{{ExecutionDate}}": human_date(st.session_state.ExecutionDate) if st.session_state.ExecutionDate else "",
                                    "{{AuthorizedRep}}": "",
                                    "{{JobNumber}}": st.session_state.JobNumber or "",
                                    "{{PropertyDescription}}": st.session_state.PropertyDescription or "",
                                    "{{FirstDeliveryDate}}": human_date(st.session_state.first_delivery) if st.session_state.first_delivery else ""
                                }

                                replace_placeholders_in_doc(doc, mapping)

                                # save to temp and provide download
                                out_tf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                                try:
                                    doc.save(out_tf.name)
                                    out_tf.close()
                                    display_name = f"Lienify_AZ_{ptype}_{'Conditional' if cond=='Yes' else 'Unconditional'}_{datetime.today().strftime('%Y%m%d')}.docx"
                                    with open(out_tf.name, "rb") as f:
                                        st.download_button("Download Filled Waiver (.docx)", data=f, file_name=display_name,
                                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                                    st.success("Generated document includes all provided details.")
                                finally:
                                    try:
                                        os.unlink(out_tf.name)
                                    except Exception:
                                        pass
